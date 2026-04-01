"""Construye un DOCX profesional de alcance para stakeholders.

Toma un ScopeSpec (dict) + mockups (MockupResult list) y genera un documento
con portada, imagenes embebidas, tablas estilizadas y firmas.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .mockup_generator import MockupResult


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


class ScopeDocxBuilder:
    """Builds a professional stakeholder scope DOCX from a ScopeSpec."""

    def __init__(self, spec: dict, mockups: list[MockupResult],
                 primary_color: str = "#C41E2A", secondary_color: str = "#1A1A1A"):
        self.spec = spec
        self.mockups = {m.id: m for m in mockups}
        self.mockups_list = mockups
        self.primary = primary_color
        self.secondary = secondary_color
        self.primary_rgb = _hex_to_rgb(primary_color)
        self.secondary_rgb = _hex_to_rgb(secondary_color)
        self.gray_rgb = RGBColor(0x66, 0x66, 0x66)
        self.doc = Document()
        self._setup()

    def _setup(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        style = self.doc.styles['Normal']
        style.font.name = 'Segoe UI'
        style.font.size = Pt(11)
        style.font.color.rgb = self.secondary_rgb

    def _heading(self, text: str, level: int = 1):
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = self.primary_rgb if level in (1, 3) else self.secondary_rgb
            if level == 1:
                run.font.size = Pt(24)
            elif level == 2:
                run.font.size = Pt(18)
            elif level == 3:
                run.font.size = Pt(14)

    def _body(self, text: str, bold: bool = False):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = self.secondary_rgb
        run.bold = bold
        return p

    def _bullet(self, text: str, bold_prefix: str = ""):
        p = self.doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            run = p.add_run(bold_prefix)
            run.bold = True
            run.font.size = Pt(11)
            run = p.add_run(text)
            run.font.size = Pt(11)
        else:
            run = p.add_run(text)
            run.font.size = Pt(11)

    def _quote(self, text: str, author: str = ""):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f'"{text}"')
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.gray_rgb
        if author:
            run2 = p.add_run(f"\n— {author}")
            run2.font.size = Pt(10)
            run2.font.color.rgb = self.primary_rgb
            run2.bold = True

    def _styled_table(self, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for i, hdr in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_shading(cell, self.secondary)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci >= len(headers):
                    break
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                text = val if isinstance(val, str) else str(val)
                run = p.add_run(text)
                run.font.size = Pt(9)
                if ri % 2 == 0:
                    _set_cell_shading(cell, "F5F5F5")
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    if i < len(row.cells):
                        row.cells[i].width = Cm(w)

    def _mockup(self, mockup_id: str, caption: str = ""):
        m = self.mockups.get(mockup_id)
        if not m or not os.path.exists(m.path):
            return
        self.doc.add_picture(m.path, width=Inches(6.5))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = self.gray_rgb

    def build(self, output_path: str) -> str:
        """Build the complete DOCX and save it."""
        self._cover_page()
        self.doc.add_page_break()
        self._table_of_contents()
        self.doc.add_page_break()
        self._executive_summary()
        self.doc.add_page_break()
        self._before_after()
        self.doc.add_page_break()
        self._role_changes()
        self.doc.add_page_break()
        self._process_flow()
        self.doc.add_page_break()
        self._screen_mockups()
        self.doc.add_page_break()
        self._worked_examples()
        self.doc.add_page_break()
        self._state_lifecycle()
        self._integrations()
        self.doc.add_page_break()
        self._roles_permissions()
        self.doc.add_page_break()
        self._exclusions()
        self._phases()
        self.doc.add_page_break()
        self._investment()
        self._faq()
        self.doc.add_page_break()
        self._glossary()
        self.doc.add_page_break()
        self._signatures()
        self.doc.add_page_break()
        self._annexes()

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path

    def _cover_page(self):
        proj = self.spec.get("project", {})
        for _ in range(4):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(proj.get("title", "Documento de Alcance"))
        run.font.size = Pt(36)
        run.font.color.rgb = self.primary_rgb
        run.bold = True
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(proj.get("subtitle", ""))
        run.font.size = Pt(18)
        run.font.color.rgb = self.secondary_rgb
        self.doc.add_paragraph()
        for text in [proj.get("company_name", ""), f"Version {proj.get('version', '1.0')} — {proj.get('date', '')}",
                     f"Preparado por: {proj.get('prepared_by', '')}"]:
            if text:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.color.rgb = self.gray_rgb
        # Approval table
        for _ in range(3):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APROBACIONES REQUERIDAS")
        run.font.size = Pt(12)
        run.bold = True
        sigs = self.spec.get("signatures", [])
        if sigs:
            self._styled_table(
                ["Nombre", "Cargo", "Firma", "Fecha"],
                [[s.get("name", ""), s.get("title", ""), "_______________", "___/___/____"] for s in sigs],
                col_widths=[5, 5.5, 3, 2.5],
            )
        notice = proj.get("confidentiality_notice", "")
        if notice:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"\n{notice}")
            run.font.size = Pt(8)
            run.italic = True
            run.font.color.rgb = self.gray_rgb

    def _table_of_contents(self):
        self._heading("Indice de Contenidos", 1)
        sections = [
            "1. Resumen Ejecutivo", "2. Antes y Despues", "3. Que Cambia Para Ti",
            "4. Flujo Completo", "5. Pantallas del Portal", "6. Ejemplos Completos",
            "7. Estados / Ciclo de Vida", "8. Integraciones",
            "9. Roles y Permisos", "10. Exclusiones", "11. Fases de Entrega",
            "12. Inversion Economica", "13. Preguntas Frecuentes",
            "14. Glosario", "15. Aprobacion y Firmas", "16. Anexos",
        ]
        for s in sections:
            p = self.doc.add_paragraph(s)
            p.paragraph_format.space_after = Pt(2)

    def _executive_summary(self):
        self._heading("1. Resumen Ejecutivo", 1)
        es = self.spec.get("executive_summary", {})
        if es.get("problem_description"):
            self._heading("El problema que resolvemos", 2)
            self._body(es["problem_description"])
        pq = es.get("problem_quote")
        if pq:
            self._quote(pq.get("text", ""), pq.get("author", ""))
        if es.get("solution_description"):
            self._heading("La solucion", 2)
            self._body(es["solution_description"])
        for b in es.get("solution_bullets", []):
            self._bullet(b.get("text", ""), bold_prefix=b.get("bold_prefix", ""))
        sq = es.get("solution_quote")
        if sq:
            self._quote(sq.get("text", ""), sq.get("author", ""))
        unchanged = es.get("unchanged_items", [])
        if unchanged:
            self._heading("Lo que NO cambia", 2)
            for b in unchanged:
                self._bullet(b.get("text", ""), bold_prefix=b.get("bold_prefix", ""))

    def _before_after(self):
        self._heading("2. Antes y Despues — Comparativa Visual", 1)
        self._mockup("before_after", "Comparativa del proceso actual vs. el nuevo sistema")

    def _role_changes(self):
        self._heading("3. Que Cambia Para Ti — Por Rol", 1)
        for sh in self.spec.get("stakeholders", []):
            self._heading(f"{sh.get('name', '')} — {sh.get('title', '')}", 2)
            self._body(f"Rol: {sh.get('role', '')}")

    def _process_flow(self):
        self._heading("4. Flujo Completo de una Operacion", 1)
        self._mockup("process_flow", "Flujo completo de una operacion")

    def _screen_mockups(self):
        self._heading("5. Pantallas del Portal — Mockups Visuales", 1)
        self._body("A continuacion se muestran las pantallas principales. Todos los datos son reales o representativos.")
        for screen in self.spec.get("screens", []):
            sid = screen.get("id", "")
            self._heading(f"{screen.get('title', sid)}", 2)
            if screen.get("description"):
                self._body(screen["description"])
            self._mockup(sid, f"Mockup: {screen.get('title', '')}")

    def _worked_examples(self):
        self._heading("6. Ejemplos Completos", 1)
        for i, ex in enumerate(self.spec.get("examples", []), 1):
            self._heading(f"Ejemplo {i}: {ex.get('title', '')}", 2)
            if ex.get("subtitle"):
                self._body(ex["subtitle"])
            dt = ex.get("data_table")
            if dt and dt.get("headers"):
                rows = [[str(c) if isinstance(c, str) else str(c) for c in row] for row in dt.get("rows", [])]
                self._styled_table(dt["headers"], rows)
            tl = ex.get("timeline")
            if tl and tl.get("headers"):
                rows = [[str(c) for c in row] for row in tl.get("rows", [])]
                self._styled_table(tl["headers"], rows)
            docs = ex.get("documents", [])
            if docs:
                self._heading("Documentacion", 3)
                for d in docs:
                    self._bullet(d)
            q = ex.get("quote")
            if q:
                self._quote(q.get("text", ""), q.get("author", ""))

    def _state_lifecycle(self):
        self._heading("7. Estados de una Operacion — Ciclo de Vida", 1)
        self._mockup("state_lifecycle", "Ciclo de vida con estados y transiciones")
        sl = self.spec.get("state_lifecycle", {})
        states = sl.get("states", [])
        if states:
            self._styled_table(
                ["Estado", "Quien lo dispara", "Que significa"],
                [[s.get("name", ""), s.get("triggered_by", ""), s.get("description", "")] for s in states],
            )

    def _integrations(self):
        self._heading("8. Integraciones — Que Se Conecta Con Que", 1)
        self._mockup("integrations", "Diagrama de integraciones entre sistemas")
        intg = self.spec.get("integrations", {})
        for fn in intg.get("footnotes", []):
            self._bullet(fn)

    def _roles_permissions(self):
        self._heading("9. Roles y Permisos", 1)
        self._mockup("roles_matrix", "Matriz de permisos por rol")

    def _exclusions(self):
        self._heading("10. Que NO Incluye Este Proyecto", 1)
        excl = self.spec.get("exclusions", [])
        if excl:
            self._styled_table(
                ["Exclusion", "Motivo"],
                [[e.get("item", ""), e.get("reason", "")] for e in excl],
            )

    def _phases(self):
        self._heading("11. Fases de Entrega y Calendario", 1)
        self._mockup("phases", "Cronograma de fases de entrega")
        for phase in self.spec.get("phases", []):
            self._heading(f"{phase.get('name', '')} ({phase.get('duration', '')})", 3)
            for d in phase.get("deliverables", []):
                self._bullet(d)

    def _investment(self):
        self._heading("12. Inversion Economica", 1)
        inv = self.spec.get("investment", [])
        if inv:
            self._styled_table(
                ["Concepto", "Coste", "Tipo"],
                [[i.get("concept", ""), i.get("cost", ""), i.get("cost_type", "")] for i in inv],
            )

    def _faq(self):
        self._heading("13. Preguntas Frecuentes", 1)
        for faq in self.spec.get("faq", []):
            self._body(f"P: {faq.get('question', '')}", bold=True)
            self._body(f"R: {faq.get('answer', '')}")

    def _glossary(self):
        self._heading("14. Glosario", 1)
        terms = self.spec.get("glossary", [])
        if terms:
            self._styled_table(
                ["Termino", "Definicion"],
                [[t.get("term", ""), t.get("definition", "")] for t in terms],
            )

    def _signatures(self):
        self._heading("15. Aprobacion y Firmas", 1)
        self._body("He revisado este documento de alcance y apruebo las funcionalidades, pantallas, integraciones, fases de entrega y exclusiones aqui descritas.")
        for sig in self.spec.get("signatures", []):
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run(sig.get("name", ""))
            run.bold = True
            run.font.size = Pt(12)
            p = self.doc.add_paragraph()
            run = p.add_run(sig.get("title", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = self.gray_rgb
            self._body("Firma: _____________________________     Fecha: ____/____/____")

    def _annexes(self):
        self._heading("16. Anexos — Referencia Visual Completa", 1)
        self._body("Todas las pantallas y diagramas generados como referencia para el desarrollo.")
        for m in self.mockups_list:
            self._body(m.caption or m.id, bold=True)
            if os.path.exists(m.path):
                self.doc.add_picture(m.path, width=Inches(6))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph()

        # Footer
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("— Fin del documento —")
        run.font.size = Pt(10)
        run.font.color.rgb = self.gray_rgb
        run.italic = True


def build_scope_docx(
    spec: dict,
    mockups: list[MockupResult],
    primary_color: str = "#C41E2A",
    secondary_color: str = "#1A1A1A",
    output_path: str = "output/alcance_stakeholders.docx",
) -> str:
    """Build the professional stakeholder scope DOCX."""
    builder = ScopeDocxBuilder(spec, mockups, primary_color, secondary_color)
    return builder.build(output_path)
