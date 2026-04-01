"""Genera documentos de salida en Markdown y DOCX."""

import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_markdown(content: str, output_path: str) -> str:
    """Guarda contenido como Markdown."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def save_docx(content: str, output_path: str, title: str = "Documento de Requerimientos") -> str:
    """Convierte Markdown a DOCX con formato profesional."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    in_table = False
    table_rows: list[list[str]] = []

    for line in content.split("\n"):
        stripped = line.strip()

        # Detectar filas de tabla
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Ignorar separadores de tabla (|---|---|)
            if all(re.match(r"^-+$", c) or c == "" for c in cells):
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        elif in_table:
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        if not stripped:
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(_clean_md(stripped[2:]), level=1)
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            doc.add_heading(_clean_md(stripped[3:]), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(_clean_md(stripped[4:]), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(_clean_md(stripped[5:]), level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(_clean_md(stripped[2:]), style="List Bullet")
        elif stripped.startswith("---"):
            doc.add_page_break()
        elif stripped.startswith("```"):
            continue
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)

    if in_table:
        _flush_table(doc, table_rows)

    doc.save(output_path)
    return output_path


def _flush_table(doc: Document, rows: list[list[str]]):
    """Inserta una tabla en el documento DOCX."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols, style="Light Grid Accent 1")
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                table.rows[i].cells[j].text = _clean_md(cell_text)


def _clean_md(text: str) -> str:
    """Elimina marcadores Markdown del texto."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*(.*?)\*", r"\1", text)  # italic
    text = re.sub(r"`(.*?)`", r"\1", text)  # code
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)  # links
    return text.strip()


def _add_rich_text(paragraph, text: str):
    """Agrega texto con formato básico (bold) a un párrafo."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(_clean_md(part))
            run.bold = True
        else:
            paragraph.add_run(_clean_md(part))
