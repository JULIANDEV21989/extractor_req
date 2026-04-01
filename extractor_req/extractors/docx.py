"""Extrae texto estructurado de archivos DOCX."""

from docx import Document
from docx.table import Table


def _table_to_markdown(table: Table) -> str:
    """Convierte una tabla DOCX a Markdown."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        num_cols = len(table.rows[0].cells)
        rows.insert(1, "| " + " | ".join(["---"] * num_cols) + " |")
    return "\n".join(rows)


def extract_docx(file_path: str) -> str:
    """Extrae contenido de un DOCX preservando estructura, estilos y tablas."""
    doc = Document(file_path)
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if not text:
                        continue
                    style = para.style.name if para.style else ""
                    if "Heading 1" in style:
                        parts.append(f"# {text}")
                    elif "Heading 2" in style:
                        parts.append(f"## {text}")
                    elif "Heading 3" in style:
                        parts.append(f"### {text}")
                    elif "List" in style or "Bullet" in style:
                        parts.append(f"- {text}")
                    else:
                        parts.append(text)
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append(_table_to_markdown(table))
                    break

    return "\n\n".join(parts) if parts else "[DOCX sin contenido extraíble]"
